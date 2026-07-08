"""
genera_factsheet_word.py
Funzione riutilizzabile: riceve i dati estratti da Claude e produce un .docx Lumen.
Valori N/D vengono evidenziati in giallo per facilitare la compilazione manuale.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


# ── Palette Lumen ────────────────────────────────────────────────────────────
VERDE_NOTTE   = RGBColor(0x0a, 0x2e, 0x22)
VERDE_SMERALDO = RGBColor(0x0F, 0x6E, 0x56)
ORO           = RGBColor(0xC9, 0xA8, 0x4C)
BIANCO        = RGBColor(0xFF, 0xFF, 0xFF)
GRIGIO        = RGBColor(0x55, 0x55, 0x55)
GIALLO_ND     = RGBColor(0xFF, 0xF3, 0xA0)  # evidenziazione valori N/D


def _set_cell_bg(cell, hex_color: str):
    """Imposta il colore di sfondo di una cella."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_para_border_bottom(para, color_hex: str, size: int = 6):
    """Aggiunge un bordo inferiore a un paragrafo."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(para, text: str, bold=False, italic=False, color=None, size_pt=10, font="Montserrat"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _highlight_if_nd(run, value: str):
    """Evidenzia in giallo se il valore è N/D."""
    if value in ("N/D", "N/A", ""):
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)


def _section_label(doc, text: str):
    """Intestazione di sezione con bordo oro."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    _set_para_border_bottom(para, "C9A84C", 6)
    run = para.add_run(text.upper())
    run.font.name = "Montserrat"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = VERDE_SMERALDO
    return para


def _numero_card(table, col_idx, valore, label, sub_label=""):
    """Popola una cella della tabella numeri chiave."""
    cell = table.cell(0, col_idx)
    _set_cell_bg(cell, "0a2e22")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_val = cell.paragraphs[0]
    p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_val.paragraph_format.space_before = Pt(6)
    run_val = p_val.add_run(valore)
    run_val.font.name = "Cormorant Garamond"
    run_val.font.size = Pt(22)
    run_val.font.bold = True
    run_val.font.color.rgb = ORO
    _highlight_if_nd(run_val, valore)

    p_lbl = cell.add_paragraph()
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_lbl = p_lbl.add_run(label.upper())
    run_lbl.font.name = "Montserrat"
    run_lbl.font.size = Pt(7)
    run_lbl.font.bold = True
    run_lbl.font.color.rgb = BIANCO

    if sub_label:
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(6)
        run_sub = p_sub.add_run(sub_label)
        run_sub.font.name = "Montserrat"
        run_sub.font.size = Pt(7)
        run_sub.font.color.rgb = ORO


def genera_factsheet_word(dati: dict, titolo: str, fonte: str, scadenza: str) -> bytes:
    """
    Genera un factsheet Word (.docx) a partire dai dati estratti da Claude.
    Restituisce i bytes del file .docx.
    """
    doc = Document()

    # ── Impostazioni pagina A4 ────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

    # ── HEADER ───────────────────────────────────────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.style = "Table Grid"
    header_tbl.autofit = False
    header_tbl.columns[0].width = Cm(12)
    header_tbl.columns[1].width = Cm(7)

    # Cella sinistra — logo testuale
    cell_logo = header_tbl.cell(0, 0)
    _set_cell_bg(cell_logo, "0a2e22")
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_logo.paragraphs[0]
    p_logo.paragraph_format.space_before = Pt(6)
    p_logo.paragraph_format.space_after = Pt(2)
    run_logo = p_logo.add_run("◆  LUMEN")
    run_logo.font.name = "Montserrat"
    run_logo.font.size = Pt(14)
    run_logo.font.bold = True
    run_logo.font.color.rgb = BIANCO
    p_sub = cell_logo.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(6)
    run_sub = p_sub.add_run("ADVISORS  ·  ADVISORY · PLANNING · WEALTH")
    run_sub.font.name = "Montserrat"
    run_sub.font.size = Pt(7)
    run_sub.font.color.rgb = ORO

    # Cella destra — badge
    cell_badge = header_tbl.cell(0, 1)
    _set_cell_bg(cell_badge, "0a2e22")
    cell_badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_badge = cell_badge.paragraphs[0]
    p_badge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_badge.paragraph_format.space_before = Pt(6)
    p_badge.paragraph_format.space_after = Pt(6)
    run_badge = p_badge.add_run("OPPORTUNITÀ DI FINANZIAMENTO")
    run_badge.font.name = "Montserrat"
    run_badge.font.size = Pt(8)
    run_badge.font.bold = True
    run_badge.font.color.rgb = ORO

    # ── HERO ─────────────────────────────────────────────────────────────
    hero_tbl = doc.add_table(rows=1, cols=1)
    hero_tbl.style = "Table Grid"
    hero_tbl.autofit = False
    hero_tbl.columns[0].width = Cm(19)

    cell_hero = hero_tbl.cell(0, 0)
    _set_cell_bg(cell_hero, "0a2e22")

    p_eyebrow = cell_hero.paragraphs[0]
    p_eyebrow.paragraph_format.space_before = Pt(4)
    run_ey = p_eyebrow.add_run(dati.get("ente_promotore", fonte).upper())
    run_ey.font.name = "Montserrat"
    run_ey.font.size = Pt(8)
    run_ey.font.bold = True
    run_ey.font.color.rgb = ORO

    p_title = cell_hero.add_paragraph()
    run_title = p_title.add_run(titolo.upper())
    run_title.font.name = "Montserrat"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = BIANCO

    p_sub_hero = cell_hero.add_paragraph()
    p_sub_hero.paragraph_format.space_after = Pt(6)
    run_sub_hero = p_sub_hero.add_run(dati.get("sottotitolo", ""))
    run_sub_hero.font.name = "Montserrat"
    run_sub_hero.font.size = Pt(10)
    run_sub_hero.font.italic = True
    run_sub_hero.font.color.rgb = ORO

    # ── INTRO BOX ────────────────────────────────────────────────────────
    doc.add_paragraph()
    intro_tbl = doc.add_table(rows=1, cols=1)
    intro_tbl.style = "Table Grid"
    intro_tbl.autofit = False
    intro_tbl.columns[0].width = Cm(19)
    cell_intro = intro_tbl.cell(0, 0)
    _set_cell_bg(cell_intro, "F7F7F5")
    p_intro = cell_intro.paragraphs[0]
    p_intro.paragraph_format.space_before = Pt(4)
    p_intro.paragraph_format.space_after = Pt(4)
    run_intro = p_intro.add_run(dati.get("descrizione_intro", ""))
    run_intro.font.name = "Montserrat"
    run_intro.font.size = Pt(10)

    # ── NUMERI CHIAVE ────────────────────────────────────────────────────
    doc.add_paragraph()
    _section_label(doc, "Numeri Chiave")

    num_tbl = doc.add_table(rows=1, cols=3)
    num_tbl.style = "Table Grid"
    num_tbl.autofit = False
    col_w = Cm(6.2)
    for col in num_tbl.columns:
        col.width = col_w

    dotazione = dati.get("dotazione_totale", "N/D")
    agev_princ = dati.get("agevolazione_principale", "N/D")
    invest_range = dati.get("investimento_range", "N/D")

    _numero_card(num_tbl, 0, dotazione, "Dotazione Totale")
    _numero_card(num_tbl, 1, agev_princ, "Agevolazione Principale", dati.get("agevolazione_nota", ""))
    _numero_card(num_tbl, 2, invest_range, "Investimento Ammissibile", dati.get("investimento_nota", ""))

    # ── STRUTTURA AGEVOLAZIONE ───────────────────────────────────────────
    doc.add_paragraph()
    _section_label(doc, "Struttura dell'Agevolazione")

    inv_tbl = doc.add_table(rows=1, cols=2)
    inv_tbl.style = "Table Grid"
    inv_tbl.autofit = False
    inv_tbl.columns[0].width = Cm(9.5)
    inv_tbl.columns[1].width = Cm(9.5)

    for col_idx, (lbl, val, nota) in enumerate([
        ("Investimento Minimo", dati.get("investimento_minimo", "N/D"), dati.get("investimento_minimo_nota", "")),
        ("Investimento Massimo", dati.get("investimento_massimo", "N/D"), dati.get("investimento_massimo_nota", "")),
    ]):
        cell = inv_tbl.cell(0, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_lbl = cell.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(4)
        run_lbl = p_lbl.add_run(lbl.upper())
        run_lbl.font.name = "Montserrat"
        run_lbl.font.size = Pt(8)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = VERDE_SMERALDO
        p_val = cell.add_paragraph()
        run_val = p_val.add_run(val)
        run_val.font.name = "Montserrat"
        run_val.font.size = Pt(16)
        run_val.font.bold = True
        run_val.font.color.rgb = VERDE_NOTTE
        _highlight_if_nd(run_val, val)
        if nota:
            p_nota = cell.add_paragraph()
            p_nota.paragraph_format.space_after = Pt(4)
            run_nota = p_nota.add_run(nota)
            run_nota.font.name = "Montserrat"
            run_nota.font.size = Pt(8)
            run_nota.font.color.rgb = GRIGIO

    # Mix agevolazione (se presente)
    mix_fp = dati.get("mix_fp_pct", 0)
    mix_fin = dati.get("mix_fin_pct", 0)
    if mix_fp > 0 and mix_fin > 0:
        p_mix = doc.add_paragraph()
        p_mix.paragraph_format.space_before = Pt(6)
        run_mix = p_mix.add_run(f"Composizione agevolazione: {dati.get('mix_fp_label','')}  +  {dati.get('mix_fin_label','')}")
        run_mix.font.name = "Montserrat"
        run_mix.font.size = Pt(9)
        run_mix.font.color.rgb = GRIGIO
        if dati.get("mix_nota"):
            p_mix_nota = doc.add_paragraph()
            run_mn = p_mix_nota.add_run(dati.get("mix_nota", ""))
            run_mn.font.name = "Montserrat"
            run_mn.font.size = Pt(8)
            run_mn.font.color.rgb = GRIGIO

    # Note aggiuntive
    note = dati.get("note_aggiuntive", "")
    if note and note not in ("N/D", ""):
        note_tbl = doc.add_table(rows=1, cols=1)
        note_tbl.style = "Table Grid"
        note_tbl.autofit = False
        note_tbl.columns[0].width = Cm(19)
        cell_note = note_tbl.cell(0, 0)
        _set_cell_bg(cell_note, "EAF4F0")
        p_note = cell_note.paragraphs[0]
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after = Pt(4)
        run_note_lbl = p_note.add_run("Note:  ")
        run_note_lbl.font.name = "Montserrat"
        run_note_lbl.font.size = Pt(9)
        run_note_lbl.font.bold = True
        run_note_lbl.font.color.rgb = VERDE_NOTTE
        run_note = p_note.add_run(note)
        run_note.font.name = "Montserrat"
        run_note.font.size = Pt(9)

    # ── BENEFICIARI E TEMPISTICHE ────────────────────────────────────────
    doc.add_paragraph()
    _section_label(doc, "Beneficiari e Tempistiche")

    bt_tbl = doc.add_table(rows=1, cols=2)
    bt_tbl.style = "Table Grid"
    bt_tbl.autofit = False
    bt_tbl.columns[0].width = Cm(9.5)
    bt_tbl.columns[1].width = Cm(9.5)

    # Colonna sinistra — Chi può candidarsi
    cell_chi = bt_tbl.cell(0, 0)
    p_chi_title = cell_chi.paragraphs[0]
    p_chi_title.paragraph_format.space_before = Pt(4)
    _set_para_border_bottom(p_chi_title, "C9A84C", 4)
    run_chi_t = p_chi_title.add_run("CHI PUÒ CANDIDARSI")
    run_chi_t.font.name = "Montserrat"
    run_chi_t.font.size = Pt(8)
    run_chi_t.font.bold = True
    run_chi_t.font.color.rgb = VERDE_NOTTE

    for req in dati.get("chi_candidarsi", []):
        p_req = cell_chi.add_paragraph(style="List Bullet")
        p_req.paragraph_format.left_indent = Cm(0.5)
        run_req = p_req.add_run(req)
        run_req.font.name = "Montserrat"
        run_req.font.size = Pt(9)

    # Colonna destra — Tempistiche
    cell_temp = bt_tbl.cell(0, 1)
    p_temp_title = cell_temp.paragraphs[0]
    p_temp_title.paragraph_format.space_before = Pt(4)
    _set_para_border_bottom(p_temp_title, "C9A84C", 4)
    run_temp_t = p_temp_title.add_run("TEMPISTICHE PRINCIPALI")
    run_temp_t.font.name = "Montserrat"
    run_temp_t.font.size = Pt(8)
    run_temp_t.font.bold = True
    run_temp_t.font.color.rgb = VERDE_NOTTE

    for t in dati.get("tempistiche", []):
        p_t = cell_temp.add_paragraph()
        p_t.paragraph_format.space_before = Pt(2)
        run_t_lbl = p_t.add_run(f"{t.get('label','')}: ")
        run_t_lbl.font.name = "Montserrat"
        run_t_lbl.font.size = Pt(9)
        run_t_lbl.font.bold = True
        run_t_lbl.font.color.rgb = VERDE_NOTTE
        run_t_val = p_t.add_run(t.get("value", "N/D"))
        run_t_val.font.name = "Montserrat"
        run_t_val.font.size = Pt(9)
        run_t_val.font.color.rgb = GRIGIO
        _highlight_if_nd(run_t_val, t.get("value", "N/D"))

    # ── CTA FOOTER ───────────────────────────────────────────────────────
    doc.add_paragraph()
    cta_tbl = doc.add_table(rows=1, cols=2)
    cta_tbl.style = "Table Grid"
    cta_tbl.autofit = False
    cta_tbl.columns[0].width = Cm(13)
    cta_tbl.columns[1].width = Cm(6)

    cell_cta_l = cta_tbl.cell(0, 0)
    _set_cell_bg(cell_cta_l, "0a2e22")
    cell_cta_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_cta_t = cell_cta_l.paragraphs[0]
    p_cta_t.paragraph_format.space_before = Pt(6)
    run_cta_t = p_cta_t.add_run("Lumen Advisors supporta la vostra candidatura")
    run_cta_t.font.name = "Montserrat"
    run_cta_t.font.size = Pt(11)
    run_cta_t.font.bold = True
    run_cta_t.font.color.rgb = BIANCO
    p_cta_d = cell_cta_l.add_paragraph()
    p_cta_d.paragraph_format.space_after = Pt(6)
    run_cta_d = p_cta_d.add_run("Verifica di eligibilità  ·  Strutturazione del programma  ·  Predisposizione della domanda  ·  Monitoraggio e rendicontazione")
    run_cta_d.font.name = "Montserrat"
    run_cta_d.font.size = Pt(8)
    run_cta_d.font.color.rgb = ORO

    cell_cta_r = cta_tbl.cell(0, 1)
    _set_cell_bg(cell_cta_r, "0a2e22")
    cell_cta_r.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for contact in ["info@lumenadvisors.it", "+41 79 601 5800", "www.lumenadvisors.it"]:
        p_c = cell_cta_r.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_c = p_c.add_run(contact)
        run_c.font.name = "Montserrat"
        run_c.font.size = Pt(9)
        run_c.font.color.rgb = ORO

    # ── PAGE FOOTER ──────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_tbl = doc.add_table(rows=1, cols=3)
    footer_tbl.style = "Table Grid"
    footer_tbl.autofit = False
    footer_tbl.columns[0].width = Cm(7)
    footer_tbl.columns[1].width = Cm(8)
    footer_tbl.columns[2].width = Cm(4)

    cell_fl = footer_tbl.cell(0, 0)
    p_fl = cell_fl.paragraphs[0]
    run_fl = p_fl.add_run("◆  LUMEN ADVISORS")
    run_fl.font.name = "Montserrat"
    run_fl.font.size = Pt(8)
    run_fl.font.bold = True
    run_fl.font.color.rgb = VERDE_NOTTE
    p_fl2 = cell_fl.add_paragraph()
    run_fl2 = p_fl2.add_run("ADVISORY · PLANNING · WEALTH")
    run_fl2.font.name = "Montserrat"
    run_fl2.font.size = Pt(7)
    run_fl2.font.color.rgb = GRIGIO

    cell_fc = footer_tbl.cell(0, 1)
    p_fc = cell_fc.paragraphs[0]
    p_fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fc = p_fc.add_run("Documento riservato  ·  Uso esclusivo del destinatario")
    run_fc.font.name = "Montserrat"
    run_fc.font.size = Pt(8)
    run_fc.font.italic = True
    run_fc.font.color.rgb = GRIGIO

    cell_fr = footer_tbl.cell(0, 2)
    p_fr = cell_fr.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_fr = p_fr.add_run("www.lumenadvisors.it")
    run_fr.font.name = "Montserrat"
    run_fr.font.size = Pt(8)
    run_fr.font.color.rgb = VERDE_SMERALDO

    # ── Salva in memoria ─────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
